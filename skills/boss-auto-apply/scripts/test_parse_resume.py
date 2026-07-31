# test_parse_resume.py
import json, os, subprocess, sys, tempfile
import unittest

from parse_resume import parse_docx, parse_pdf, extract_fields

class TestParseResume(unittest.TestCase):
    def setUp(self):
        # 创建最小测试 docx
        self.tmpdir = tempfile.mkdtemp()
        self.docx_path = os.path.join(self.tmpdir, "resume.docx")
        from docx import Document
        doc = Document()
        doc.add_paragraph("张三")
        doc.add_paragraph("3年Python后端开发经验，熟悉Django、FastAPI")
        doc.add_paragraph("期望岗位：Python后端开发")
        doc.add_paragraph("期望城市：北京")
        doc.add_paragraph("期望薪资：20-30K")
        doc.save(self.docx_path)

    def test_parse_docx(self):
        text = parse_docx(self.docx_path)
        self.assertIn("张三", text)
        self.assertIn("Python", text)

    def test_extract_fields(self):
        text = "张三\n3年Python后端开发经验，熟悉Django、FastAPI\n期望岗位：Python后端开发\n期望城市：北京\n期望薪资：20-30K"
        fields = extract_fields(text)
        self.assertEqual(fields["name"], "张三")
        self.assertIn("Python", fields["skills"])
        self.assertIn("后端", fields["expected_job"])
        self.assertEqual(fields["city"], "北京")
        self.assertIn("20-30K", fields["expected_salary"])

    def test_cli_outputs_markdown(self):
        # 通过 CLI 调用验证输出 Markdown 且包含字段
        out_path = os.path.join(self.tmpdir, "out.md")
        result = subprocess.run(
            [sys.executable, "parse_resume.py", self.docx_path, "--output", out_path],
            capture_output=True, text=True, cwd=os.path.dirname(__file__),
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        with open(out_path, encoding="utf-8") as f:
            content = f.read()
        self.assertIn("张三", content)
        self.assertIn("姓名", content)

if __name__ == "__main__":
    unittest.main()
